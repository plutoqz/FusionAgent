from __future__ import annotations

from pathlib import Path

import geopandas as gpd
from shapely.geometry import box
from docx import Document

from scripts.clean_benin_final_buildings import (
    CandidateNameRule,
    _apply_shp_cleaning,
    _candidate_decision_summary,
    write_report,
    build_clean_final_height_series,
    build_final_height_value,
    extract_keep_fields_from_docx,
    load_candidate_name_rules,
    score_candidate_features_from_gpkg,
    score_candidate_non_building,
    shp_source_columns_for_keep_fields,
    should_drop_non_building_record,
)


def test_extract_keep_fields_from_docx_reads_list_in_order(tmp_path: Path) -> None:
    docx_path = tmp_path / "keep_fields.docx"
    doc = Document()
    doc.add_paragraph("01. detail_side")
    doc.add_paragraph("02. source_layer")
    doc.add_paragraph("03. group_score")
    doc.add_paragraph("04. src_flag")
    doc.add_paragraph("05. name_fused")
    doc.add_paragraph("06. name_candidates")
    doc.add_paragraph("07. fusion_lineage")
    doc.add_paragraph("08. height_final")
    doc.add_paragraph("09. height_final_source")
    doc.add_paragraph("10. fusion_source")
    doc.add_paragraph("11. final_height")
    doc.save(docx_path)

    assert extract_keep_fields_from_docx(docx_path) == [
        "detail_side",
        "source_layer",
        "group_score",
        "src_flag",
        "name_fused",
        "name_candidates",
        "fusion_lineage",
        "fusion_source",
        "final_height",
    ]


def test_non_building_record_rule_matches_clear_false_positives() -> None:
    assert should_drop_non_building_record("Clôture du cimetière", 1.5)
    assert should_drop_non_building_record("Parking Moto", 2.0)
    assert should_drop_non_building_record("Site Pompe", 0.5)
    assert should_drop_non_building_record("Funerarium Les anges", 0.5)
    assert should_drop_non_building_record("Place vodùn Yedomin", 1.0)
    assert should_drop_non_building_record("Buanderie", 1.5)
    assert should_drop_non_building_record("Local cuisson", 2.0)
    assert should_drop_non_building_record("Séchoir à riz de Sinsinkou-Tora", 1.0)
    assert should_drop_non_building_record("Abris famille, Pédiatrie", 2.0)
    assert should_drop_non_building_record("Paillote", 1.5)
    assert not should_drop_non_building_record("Mairie de Ouidah", 1.5)
    assert not should_drop_non_building_record("Basilique de l'Immaculée Conception", 1.5)
    assert not should_drop_non_building_record("Guérite", 1.0)
    assert not should_drop_non_building_record("Scierie", 1.5)
    assert not should_drop_non_building_record("Dortoir", 1.5)
    assert not should_drop_non_building_record("Marché Arzeke de Parakou", 1.5)


def test_score_candidate_non_building_drops_explicit_non_building_names_even_if_not_low() -> None:
    parking_rule = CandidateNameRule(
        name="Parking Moto",
        name_zh="摩托车停车场",
        suggested_action="建议删除",
        risk_category="停车/车库",
        suggested_non_building_category="停车/场地",
        reason="明显属于停车场或停车区域，不应视作建筑物",
    )
    reservoir_rule = CandidateNameRule(
        name="Réservoir d'Eau",
        name_zh="蓄水池/水箱",
        suggested_action="待人工判断",
        risk_category="未分类",
        suggested_non_building_category="设备/点位",
        reason="名称本身不足以判断，建议结合高度、面积和几何形态复核",
    )

    parking = score_candidate_non_building(
        name="Parking Moto",
        rule=parking_rule,
        raw_height=12.5,
        area_m2=184.05,
        aspect_ratio=2.41,
    )
    reservoir = score_candidate_non_building(
        name="Réservoir d'Eau",
        rule=reservoir_rule,
        raw_height=9.0,
        area_m2=241.11,
        aspect_ratio=1.01,
    )

    assert parking.should_drop
    assert parking.total_points >= parking.delete_threshold
    assert parking.name_points >= 5
    assert reservoir.should_drop
    assert reservoir.total_points >= reservoir.delete_threshold
    assert reservoir.name_points >= 5


def test_score_candidate_non_building_keeps_large_high_ambiguous_site_but_drops_low_small_one() -> None:
    place_rule = CandidateNameRule(
        name="Place vodùn Yedomin",
        name_zh="Yedomin 伏都场地",
        suggested_action="待人工判断",
        risk_category="未分类",
        suggested_non_building_category="停车/场地",
        reason="名称本身不足以判断，建议结合高度、面积和几何形态复核",
    )

    large_high = score_candidate_non_building(
        name="Place vodùn Yedomin",
        rule=place_rule,
        raw_height=4.5,
        area_m2=210.9,
        aspect_ratio=1.66,
    )
    low_small = score_candidate_non_building(
        name="Place vodùn Yedomin",
        rule=place_rule,
        raw_height=1.0,
        area_m2=22.1,
        aspect_ratio=1.12,
    )

    assert not large_high.should_drop
    assert large_high.total_points < large_high.delete_threshold
    assert low_small.should_drop
    assert low_small.total_points >= low_small.delete_threshold


def test_load_candidate_name_rules_reads_candidate_review_csv(tmp_path: Path) -> None:
    csv_path = tmp_path / "candidates.csv"
    csv_path.write_text(
        "name,name_zh,suggested_action,risk_category,reason,suggested_non_building_category_v2\n"
        "Parking Moto,摩托车停车场,建议删除,停车/车库,明显属于停车场或停车区域，不应视作建筑物,停车/场地\n",
        encoding="utf-8-sig",
    )

    rules = load_candidate_name_rules(csv_path)
    rule = rules["parking moto"]
    assert rule.name == "Parking Moto"
    assert rule.suggested_action == "恢复保留"
    assert rule.suggested_non_building_category == "停车/场地"


def test_apply_shp_cleaning_drops_only_requested_fids() -> None:
    frame = gpd.GeoDataFrame(
        {
            "fid": [1, 2],
            "detail_sid": ["a", "b"],
            "source_lay": ["OSM", "OSM"],
            "group_scor": [1.0, 2.0],
            "src_flag": ["OSM", "OSM"],
            "name_fused": ["Parking Moto", "Dortoir"],
            "name_candi": ["Parking Moto", "Dortoir"],
            "fusion_lin": ["x", "y"],
            "height_fin": [0.0, 0.0],
            "height_f_1": ["src", "src"],
            "fusion_sou": ["src", "src"],
            "height_c_1": [12.5, 4.0],
            "final_heig": [12.5, 4.0],
        },
        geometry=[box(0, 0, 10, 10), box(20, 0, 30, 10)],
        crs="EPSG:32631",
    )

    cleaned, stats = _apply_shp_cleaning(frame, {1})

    assert cleaned["name_fused"].tolist() == ["Dortoir"]
    assert "height_f_1" not in cleaned.columns
    assert "height_fin" not in cleaned.columns
    assert "final_heig" in cleaned.columns
    assert stats["dropped_non_building"] == 1
    assert stats["input_count"] == 2


def test_write_report_includes_chinese_categories_and_scoring_method(tmp_path: Path) -> None:
    report_path = tmp_path / "report.txt"
    candidate_summary = gpd.pd.DataFrame(
        [
            {
                "name": "Parking Moto",
                "name_zh": "摩托车停车场",
                "risk_category": "停车/车库",
                "non_building_category_zh": "停车/场地",
                "should_drop": True,
                "score": 7,
                "delete_threshold": 7,
            }
        ]
    )
    report = {
        "gpkg_input_path": "a",
        "gpkg_output_path": "b",
        "shp_input_path": "c",
        "shp_output_path": "d",
        "docx_path": "e",
        "keep_fields": ["detail_side"],
        "name_review_csv": "candidates.csv",
        "candidate_summary": candidate_summary,
        "gpkg_input_count": 10,
        "gpkg_output_count": 9,
        "gpkg_dropped_non_building": 1,
        "gpkg_floor_to_3": 0,
        "gpkg_exact_three_to_3_01": 0,
        "shp_input_count": 2,
        "shp_output_count": 2,
        "shp_dropped_non_building": 0,
        "shp_floor_to_3": 0,
        "shp_exact_three_to_3_01": 0,
    }

    write_report(report_path, report)
    content = report_path.read_text(encoding="utf-8")

    assert "打分方法" in content
    assert "删除指标" in content
    assert "Parking Moto（摩托车停车场）" in content
    assert "停车/场地" in content
    assert "总分 >= 7" in content


def test_final_height_rule_restores_between_2_5_and_3m_values() -> None:
    assert build_final_height_value(2.4) == 3.0
    assert build_final_height_value(2.5) == 2.5
    assert build_final_height_value(2.8) == 2.8
    assert build_final_height_value(3.0) == 3.01


def test_load_candidate_name_rules_only_marks_two_named_low_non_buildings_for_drop(tmp_path: Path) -> None:
    csv_path = tmp_path / "candidates.csv"
    csv_path.write_text(
        "name,name_zh,suggested_action,risk_category,reason,suggested_non_building_category_v2\n"
        "Clôture du cimetière,墓地围栏,建议删除,围栏/边界,明显属于围栏、边界或门卫类非建筑对象,边界/围挡\n"
        "Séchoir à riz de Sinsinkou-Tora,Sinsinkou-Tora 稻谷晾晒场,建议人工复核,加工/附属设施,更像生产附属设施，需结合几何和用途判断,加工附属\n"
        "Parking Moto,摩托车停车场,建议删除,停车/车库,明显属于停车场或停车区域，不应视作建筑物,停车/场地\n",
        encoding="utf-8-sig",
    )

    rules = load_candidate_name_rules(csv_path)

    assert rules["clôture du cimetière"].suggested_action == "建议删除"
    assert rules["séchoir à riz de sinsinkou-tora"].suggested_action == "建议删除"
    assert rules["parking moto"].suggested_action == "恢复保留"


def test_candidate_decision_summary_preserves_only_final_height_outputs() -> None:
    summary = _candidate_decision_summary([], {})
    assert "raw_height" in summary.columns
    assert "name_zh" in summary.columns


def test_final_height_rule_distinguishes_floor_and_exact_three() -> None:
    assert build_final_height_value(2.4) == 3.0
    assert build_final_height_value(2.6) == 2.6
    assert build_final_height_value(3.0) == 3.01
    assert build_final_height_value(8.234) == 8.23


def test_clean_final_height_series_falls_back_to_existing_final_height() -> None:
    result = build_clean_final_height_series(
        source_values=[None, 2.5, 3.0, 7.126],
        existing_final_values=[3.0, 3.0, 3.0, 9.0],
    )
    assert result.tolist() == [3.01, 2.5, 3.01, 7.13]


def test_clean_final_height_series_floors_fallback_values_too() -> None:
    result = build_clean_final_height_series(
        source_values=[None, None, None, None],
        existing_final_values=[2.2, 2.6, 3.0, 6.666],
    )
    assert result.tolist() == [3.0, 2.6, 3.01, 6.67]


def test_shp_source_columns_for_keep_fields_maps_truncated_names() -> None:
    columns = [
        "detail_sid",
        "source_lay",
        "group_scor",
        "src_flag",
        "name_fused",
        "name_candi",
        "fusion_lin",
        "fusion_sou",
        "final_heig",
    ]
    assert shp_source_columns_for_keep_fields(columns) == {
        "detail_side": "detail_sid",
        "source_layer": "source_lay",
        "group_score": "group_scor",
        "src_flag": "src_flag",
        "name_fused": "name_fused",
        "name_candidates": "name_candi",
        "fusion_lineage": "fusion_lin",
        "fusion_source": "fusion_sou",
        "final_height": "final_heig",
    }
